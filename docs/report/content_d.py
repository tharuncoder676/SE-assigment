"""Section 17: originality and similarity statement."""
import json
import pathlib

from docbuild import (bullets, code, figure, h1, h2, link_para, note, para,
                      table, FIG)
from team import MEMBERS

REPO = "https://github.com/tharuncoder676/SE-assigment"
DATA = json.loads(
    (pathlib.Path(__file__).resolve().parents[1] / "originality-report.json")
    .read_text(encoding="utf-8"))


def section17(doc):
    h1(doc, "17.  Originality and Similarity Statement", page_break=True)

    para(doc, "This section reports what we can measure about the originality of this report, "
              "states plainly what we cannot measure, and leaves space for the institutional "
              "similarity report that only the department can generate.")

    h2(doc, "17.1  Declaration")
    para(doc, "We declare that this report is our own work. The system it describes was designed "
              "and implemented by us, and every result quoted in it was produced by executing our "
              "own code. Where we have used the words of the assignment brief we have quoted and "
              "attributed them. Where we have drawn on published sources we have cited them in "
              "Section 15. Where an AI assistant was used, that use is declared in full at the "
              "end of Section 15. No part of this report has been submitted for any other "
              "assessment.")

    h2(doc, "17.2  What we measured, and how")
    para(doc, "We wrote a small analysis script, `docs/report/originality.py`, which anyone can "
              "re-run against the submitted document. It extracts every paragraph and table cell "
              "from the `.docx`, normalises the text (lower-cased, punctuation stripped), and "
              "splits it into overlapping **eight-word shingles** — eight words being the window "
              "most similarity services treat as a match. It then measures how many of those "
              "shingles also occur in the source documents we were given, and how many are "
              "repeated inside the report itself.")
    para(doc, "The measurement covers Sections 1 to 16. This section is excluded from its own "
              "figures, because a section that reports statistics about the document containing "
              "it would otherwise be counted in them.", space_after=4)
    link_para(doc, "Analysis script", REPO + "/blob/main/docs/report/originality.py")
    link_para(doc, "Raw result (JSON)", REPO + "/blob/main/docs/originality-report.json")
    link_para(doc, "Source documents compared against", REPO + "/tree/main/docs/sources")

    h2(doc, "17.3  Measured results")
    src = DATA["per_source"]
    rows = []
    for key, label in (("assignment-brief", "CSA10 assignment brief (the problem statement)"),
                       ("reference-format-sample", "Sample report supplied as a formatting guide")):
        if key in src:
            d = src[key]
            rows.append([label, "{:,}".format(d["source_words"]),
                         str(d["matched_shingles"]),
                         "**{:.2f}%**".format(d["overlap_percent"])])
    rows.append(["**Combined against all supplied sources**", "—",
                 str(DATA["matched_shingles"]),
                 "**{:.2f}%**".format(DATA["combined_source_overlap_percent"])])
    rows.append(["Internal self-similarity (repeated passages within the report)", "—", "—",
                 "**{:.2f}%**".format(DATA["internal_self_similarity_percent"])])
    table(doc,
          ["Compared against", "Source words", "Matching 8-word blocks", "Similarity"],
          rows, widths=[46, 15, 20, 19], size=8.6, align_center=(1, 2, 3))

    para(doc, "Out of **{:,} eight-word blocks** of running prose, **{}** appear in any document "
              "we were given, and the longest unbroken borrowed run is **{} words**. That run is "
              "the phrase from the brief describing what the platform must manage, which appears "
              "in Section 1.1 inside quotation marks and attributed to the brief, as a direct "
              "quotation should be.".format(DATA["prose_shingles"], DATA["matched_shingles"],
                                            DATA["longest_verbatim_run_words"]))
    para(doc, "One episode is worth recording. An earlier draft paraphrased the brief loosely "
              "and scored a lower number; rewriting the passage as an accurate, attributed "
              "quotation raised the count of matching blocks. That is the right trade — a marked, "
              "cited quotation is better scholarship than a paraphrase written to dodge a matching "
              "algorithm, even though the algorithm rewards the paraphrase. We mention it because "
              "optimising prose against a similarity score, rather than against honesty, is a "
              "habit worth naming and avoiding.")
    para(doc, "The internal self-similarity of **{:.1f}%** also deserves an explanation rather "
              "than a footnote. Section 16 contains eight independently written reflections, and "
              "every one answers the same five prompts, so the prompt headings themselves — "
              "\"Challenges I faced\", \"What I learned\", \"Course outcome attainment\" — and "
              "the course-outcome references recur eight times by design. That is the structure "
              "the assignment asks for, not padding. The bodies of the eight reflections are "
              "distinct: each member writes about the part of the system they owned."
              .format(DATA["internal_self_similarity_percent"]))

    h2(doc, "17.4  Composition of the document")
    para(doc, "Similarity tools match text without asking what kind of text it is. It is worth "
              "recording what this report is actually made of, because three of the four "
              "categories below are expected to match something and none of them is plagiarism.")
    comp = DATA["composition"]
    table(doc,
          ["Content type", "Words", "Share", "Why a similarity tool may flag it"],
          [["Original narrative prose", "{:,}".format(comp["prose"]["words"]),
            "{:.1f}%".format(comp["prose"]["percent"]),
            "It should not. This is the analysis, argument and reflection written by us"],
           ["Tables (requirements, test cases, comparisons)", "{:,}".format(comp["table"]["words"]),
            "{:.1f}%".format(comp["table"]["percent"]),
            "Short technical phrases such as HTTP status codes and quality-attribute names recur across the discipline"],
           ["Code, pseudocode and console output", "{:,}".format(comp["code"]["words"]),
            "{:.1f}%".format(comp["code"]["percent"]),
            "Our own source code and verbatim tool output. Framework keywords are identical everywhere by necessity"],
           ["Bibliography", "{:,}".format(comp["reference"]["words"]),
            "{:.1f}%".format(comp["reference"]["percent"]),
            "Citations match their sources by definition; most tools can be set to exclude the bibliography"]],
          widths=[27, 9, 8, 56], size=8.4, align_center=(1, 2))
    para(doc, "Total measured length: **{:,} words**.".format(DATA["total_words"]),
         space_after=3)

    h2(doc, "17.5  The limits of this analysis")
    note(doc, "**We want to be explicit about what this is not.** The figures above are a "
              "self-assessment against the documents we were given. They are not a Turnitin, "
              "Urkund or iThenticate result, and we make no claim to have checked this report "
              "against a global corpus of publications, web pages or previously submitted "
              "student work — we have no access to one. The definitive similarity index for this "
              "submission is the one produced by the department's own tool; Section 17.6 gives the "
              "settings we recommend for it and Section 17.8 is left for its receipt.")
    para(doc, "Two further caveats. First, a similarity score is a measure of matching text, not "
              "of academic misconduct: a correctly quoted and cited passage raises the score "
              "while being entirely proper. Second, our own measure only detects reuse of the "
              "specific sources listed in 17.3; it would not detect an unattributed source we "
              "never compared against. We state this because a number presented without its "
              "limits is misleading, and Section 12.5 commits us to not doing that.")

    h2(doc, "17.6  Institutional similarity report")
    para(doc, "The official report is to be generated by the department and attached here. When "
              "running it, we suggest the settings below, which are the ones that make a score "
              "interpretable for a software engineering report containing source code:")
    bullets(doc, [
        "**Exclude the bibliography** — Section 15 is 25 citations and will otherwise match its own sources.",
        "**Exclude quoted material** — the quotation from the brief in Section 1.1 is deliberate and attributed.",
        "**Exclude small matches** of fewer than about 8–10 consecutive words, which catch unavoidable technical phrases such as \"HTTP 409 Conflict\" or \"PBKDF2-HMAC-SHA256\".",
        "**Note that Sections 6, 7 and 9 contain code listings and verbatim console output.** These are our own, and the repository history at the link below timestamps every line of them.",
    ])
    link_para(doc, "Commit history evidencing authorship of all code", REPO + "/commits/main")

    doc.add_page_break()
    h2(doc, "17.7  Originality analysis report")
    para(doc, "The sheet below is the output of our analyser, generated from the submitted "
              "document and reproduced here in full. It carries its own provenance: the date, the "
              "script that produced it, the repository revision it was run against, and a "
              "statement on its face of what it is and is not.", space_after=4)
    figure(doc, FIG / "originality-report.png", 16.6,
           "Originality Analysis Report — generated by `docs/report/originality.py` from the "
           "submitted document. Measured similarity 0.17% against all supplied sources, with "
           "zero unattributed matches.")

    doc.add_page_break()
    h2(doc, "17.8  Attachment — institutional similarity report")
    para(doc, "The department's own similarity check is the authoritative one. Paste its receipt "
              "in the space below, then complete the record and signatures.", space_after=4)
    _placeholder(doc)

    table(doc,
          ["Field", "To be completed on submission"],
          [["Similarity index reported by the institutional tool", ""],
           ["Tool and version used", ""],
           ["Exclusions applied (bibliography / quotes / small matches)", ""],
           ["Date of check", ""],
           ["Checked by", ""]],
          widths=[46, 54], size=8.8)

    doc.add_paragraph()
    para(doc, "Signed on behalf of the team:", justify=False, space_after=8)
    table(doc,
          ["Member", "Register number", "Signature", "Date"],
          [[name, reg, "", ""] for name, reg, _ in MEMBERS],
          widths=[30, 22, 28, 20], size=8.8)


def _placeholder(doc):
    """An empty framed box for the similarity receipt."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor
    from docbuild import _borders

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    _borders(p, colour="9FB0C2", size=6)
    run = p.add_run("\n\n\n\n[  Attach the institutional similarity report / Turnitin receipt here  ]\n\n\n\n")
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x7A, 0x8A, 0x9A)
    return p
