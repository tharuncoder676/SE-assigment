"""Reusable building blocks for the assignment report.

Everything the report needs that python-docx does not expose directly (cell
shading, page-number fields, borders, keep-with-next) is done here once so the
content modules stay readable.
"""
import pathlib

from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = pathlib.Path(__file__).resolve().parents[2]
FIG = ROOT / "docs" / "figures"
SHOT = ROOT / "docs" / "screenshots"

INK = RGBColor(0x12, 0x23, 0x3A)
BRAND = RGBColor(0x0B, 0x45, 0x7C)
MUTED = RGBColor(0x5F, 0x71, 0x86)
CODE_BG = "F2F5F8"
HEAD_BG = "0B457C"
ALT_BG = "F4F8FC"

BODY_SIZE = Pt(10.2)
CONTENT_WIDTH_CM = 16.8


# ---------------------------------------------------------------- low level
def _shade(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for tag, value in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement("w:" + tag)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tcPr.append(mar)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def _borders(paragraph, colour="C7D5E2", size=6, sides=("top", "left", "bottom", "right")):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for side in sides:
        node = OxmlElement("w:" + side)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "4")
        node.set(qn("w:color"), colour)
        pbdr.append(node)
    pPr.append(pbdr)


def _field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, sep, end):
        run._r.append(node)
    return run


# ---------------------------------------------------------------- document
def new_document():
    import docx

    doc = docx.Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.9)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = BODY_SIZE
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.widow_control = True
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), "Calibri")
    return doc


def add_page_border(doc, colour="0B457C", size=8, space=22):
    """A thin rule around every page, measured from the page edge."""
    sectPr = doc.sections[0]._sectPr
    borders = OxmlElement("w:pgBorders")
    borders.set(qn("w:offsetFrom"), "page")
    for side in ("top", "left", "bottom", "right"):
        node = OxmlElement("w:" + side)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), str(space))
        node.set(qn("w:color"), colour)
        borders.append(node)
    # The schema fixes the order of sectPr children: pgBorders must follow
    # pgMar/paperSrc and precede cols. Appending blindly puts it after cols,
    # which makes Word and LibreOffice discard it silently.
    anchor = sectPr.find(qn("w:paperSrc"))
    if anchor is None:
        anchor = sectPr.find(qn("w:pgMar"))
    if anchor is not None:
        anchor.addnext(borders)
    else:
        sectPr.append(borders)


def add_header_footer(doc, title, course):
    section = doc.sections[0]

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.tab_stops.add_tab_stop(
        Cm(CONTENT_WIDTH_CM), WD_TAB_ALIGNMENT.RIGHT)
    run = header.add_run(course + "\t" + title)
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED
    _borders(header, colour="C7D5E2", size=4, sides=("bottom",))

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Page ")
    _field(footer, " PAGE ")
    run2 = footer.add_run(" of ")
    _field(footer, " NUMPAGES ")
    for r in footer.runs:
        r.font.size = Pt(8.5)
        r.font.color.rgb = MUTED


# ---------------------------------------------------------------- headings
def h1(doc, text, page_break=False):
    if page_break:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(11)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = BRAND
    _borders(p, colour="9FC0DC", size=6, sides=("bottom",))
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10.8)
    run.font.color.rgb = INK
    return p


def para(doc, text, justify=True, size=None, italic=False, space_after=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(space_after)
    _rich(p, text, size=size or BODY_SIZE, italic=italic)
    return p


def _rich(paragraph, text, size, italic=False):
    """Supports **bold**, `code` and normal text in one string."""
    import re
    for token in re.split(r"(\*\*.+?\*\*|`.+?`)", text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(float(size.pt) - 0.9)
        else:
            run = paragraph.add_run(token)
        run.font.size = size
        run.italic = italic
    return paragraph


def bullets(doc, items, size=None, numbered=False):
    for item in items:
        p = doc.add_paragraph(style="List Number" if numbered else "List Bullet")
        p.paragraph_format.space_after = Pt(1.5)
        p.paragraph_format.left_indent = Cm(0.65)
        p.paragraph_format.line_spacing = 1.06
        _rich(p, item, size=size or BODY_SIZE)


# ---------------------------------------------------------------- tables
def table(doc, headers, rows, widths, size=8.6, header_size=8.4, align_center=()):
    """widths are relative weights; they are scaled to the content width."""
    total = float(sum(widths))
    cm = [CONTENT_WIDTH_CM * w / total for w in widths]

    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False

    # Cell widths alone are advisory: Word and LibreOffice both re-flow the
    # columns unless the table declares a fixed layout and a matching grid.
    tblPr = t._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    grid = t._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        t._tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for width_cm in cm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(Cm(width_cm).twips)))
        grid.append(col)
    t._tbl.insert(list(t._tbl).index(tblPr) + 1, grid)

    for idx, head in enumerate(headers):
        cell = t.rows[0].cells[idx]
        cell.width = Cm(cm[idx])
        _shade(cell, HEAD_BG)
        _cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        run = p.add_run(head)
        run.bold = True
        run.font.size = Pt(header_size)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        if idx in align_center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for r_i, row in enumerate(rows):
        new_row = t.add_row()
        trPr = new_row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        trPr.append(cant_split)
        cells = new_row.cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            cell.width = Cm(cm[idx])
            _cell_margins(cell)
            if r_i % 2 == 1:
                _shade(cell, ALT_BG)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.line_spacing = 1.02
            _rich(p, str(value), size=Pt(size))
            if idx in align_center:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # repeat the header row when a table spans a page boundary
    trPr = t.rows[0]._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)
    return t


# ---------------------------------------------------------------- code + figures
def code(doc, text, size=7.6, caption=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.left_indent = Cm(0.15)
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = True
    _borders(p, colour="C7D5E2", size=4)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), CODE_BG)
    p._p.get_or_add_pPr().append(shading)

    lines = text.strip("\n").split("\n")
    for i, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0x1B, 0x2E, 0x44)
        if i < len(lines) - 1:
            run.add_break()
    if caption:
        cap(doc, caption)
    return p


def figure(doc, path, width_cm, caption, space_before=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    cap(doc, caption)
    return p


def cap(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(8.3)
    run.font.color.rgb = MUTED
    return p


def note(doc, text):
    """A callout used for the 'what went wrong' asides."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Cm(0.2)
    p.paragraph_format.line_spacing = 1.06
    _borders(p, colour="B4530A", size=4, sides=("left",))
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), "FFF8F0")
    p._p.get_or_add_pPr().append(shading)
    _rich(p, text, size=Pt(9.4))
    return p


def link(paragraph, text, url, size=9.6):
    """A real clickable hyperlink."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    colour = OxmlElement("w:color")
    colour.set(qn("w:val"), "0B5FAE")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Consolas")
    rFonts.set(qn("w:hAnsi"), "Consolas")
    for node in (rFonts, colour, underline, sz):
        rPr.append(node)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return paragraph


def link_para(doc, label, url, size=9.6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.2)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(label + "  ")
    run.font.size = Pt(9.6)
    link(p, url, url, size=size)
    return p
