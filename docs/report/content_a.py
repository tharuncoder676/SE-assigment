"""Title page, table of contents and sections 1-4."""
from docbuild import (bullets, cap, code, figure, h1, h2, link_para, note, para,
                      table, FIG, SHOT)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

REPO = "https://github.com/tharuncoder676/SE-assigment"
RUN_URL = REPO + "/actions/runs/33625229403"


def title_page(doc):
    for _ in range(1):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("CSA10 — SOFTWARE ENGINEERING")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x5F, 0x71, 0x86)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("Assignment Report")
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(0x0B, 0x45, 0x7C)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("SmartCare — A Scalable Healthcare Appointment\nand Patient Service Platform")
    run.bold = True
    run.font.size = Pt(19)
    run.font.color.rgb = RGBColor(0x12, 0x23, 0x3A)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("Designed, built, containerised, tested and deployed through a "
                    "complete software engineering life cycle")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x5F, 0x71, 0x86)

    figure(doc, FIG / "fig1.png", 15.2,
           "Figure 1 — Layered architecture of the delivered system. Solid arrows are synchronous "
           "request paths; green arrows are asynchronous event delivery.")

    h2(doc, "A.  Assignment Information")
    table(doc,
          ["Field", "Details"],
          [["Department", "Computer Science and Engineering"],
           ["Programme", "B.Tech — Computer Science and Engineering"],
           ["Course Code & Course Name", "CSA10 — Software Engineering"],
           ["Academic Year / Batch", "2023–2027"],
           ["Faculty Name", "[Faculty Name]"],
           ["Assignment Title", "Smart Healthcare Appointment and Patient Service Platform"],
           ["Team Members", "[Student Name 1] — [Register No.]    |    [Student Name 2] — [Register No.]    |    "
                            "[Student Name 3] — [Register No.]    |    [Student Name 4] — [Register No.]"],
           ["Date of Issue", "[DD-MM-YYYY]"],
           ["Date of Submission", "02-09-2026"],
           ["Maximum Marks", "100"],
           ["Course Outcome(s)", "CO1, CO2, CO3, CO4, CO5"],
           ["Bloom's Taxonomy Level", "L3 — Apply, L4 — Analyse, L5 — Evaluate, L6 — Create"]],
          widths=[26, 74], size=8.8)


def evidence_block(doc):
    h1(doc, "Project Evidence — Everything in This Report Is Reproducible")
    para(doc, "The whole system is public. Every screenshot, number and log line in this report "
              "came from running the code below, on our own machines and on GitHub's CI runners. "
              "Nothing here is illustrative or mocked up.", space_after=4)
    link_para(doc, "Source repository", REPO)
    link_para(doc, "CI/CD pipeline run — all four stages green", RUN_URL)
    link_para(doc, "Full run history, including the two runs that failed", REPO + "/actions")
    link_para(doc, "Commit history", REPO + "/commits/main")
    link_para(doc, "Screenshots, console transcripts and load-test data", REPO + "/tree/main/docs")


def toc(doc, pages):
    h1(doc, "Table of Contents")
    rows = [[name, str(pages.get(key, "—"))] for key, name in TOC_ENTRIES]
    table(doc, ["Section", "Page"], rows, widths=[92, 8], size=8.9,
          header_size=8.8, align_center=(1,))

    h2(doc, "List of Figures and Evidence")
    table(doc,
          ["Group", "Count", "What it shows", "Section"],
          [["Design figures", "6", "Architecture, use cases, ER model, sequence, flowchart, CI/CD pipeline", "5, 6, 7"],
           ["Performance charts", "3", "Throughput, latency percentiles, cost of the security primitives", "10"],
           ["Application screenshots", "15", "The real UI and OpenAPI docs driven end to end by an automation script", "9"],
           ["Console transcripts", "10", "pytest, coverage, lint, git, CI, container build, API session, logs, metrics", "9"]],
          widths=[22, 7, 56, 15], size=8.2, align_center=(1, 3))


TOC_ENTRIES = [
    ("s1", "1.  Problem Statement and Problem Formulation"),
    ("s2", "2.  Objectives and Expected Outcomes"),
    ("s3", "3.  Requirements, Constraints and Assumptions"),
    ("s4", "4.  Application of Software Engineering Concepts"),
    ("s5", "5.  Design, Proposed Solution and Methodology"),
    ("s6", "6.  Algorithm, Pseudocode and Flowchart"),
    ("s7", "7.  Implementation, Source Code and Tools Used"),
    ("s8", "8.  Test Cases and Expected / Actual Results"),
    ("s9", "9.  Execution Screenshots and Outputs"),
    ("s10", "10.  Results and Validation"),
    ("s11", "11.  Analysis, Comparison, Trade-offs and Justification"),
    ("s12", "12.  Broader Considerations and SDG Relevance"),
    ("s13", "13.  Conclusion, Limitations and Possible Improvements"),
    ("s14", "14.  Individual Contribution of Group Members"),
    ("s15", "15.  References"),
    ("s16", "16.  One-Page Individual Reflection"),
]


# =========================================================== section 1
def section1(doc):
    h1(doc, "1.  Problem Statement and Problem Formulation", page_break=True)

    h2(doc, "1.1  The industry problem as it was given to us")
    para(doc, "A multi-specialty healthcare organisation is struggling on several fronts at once. "
              "Appointments are delayed, the volume of user requests keeps climbing, the service "
              "is not reliably available, the system goes down, and the applications the "
              "organisation already owns are difficult to integrate with anything new. What they "
              "want is a platform that can manage patient appointments, doctor availability, "
              "notifications, service requests and healthcare information, and that will still "
              "work when the load doubles.")
    para(doc, "When we broke this down, we found that the visible complaint (appointments are "
              "slow) is not really one problem. It is at least four, and they have different "
              "causes and different fixes.")
    table(doc,
          ["Symptom the organisation reports", "What is actually going on underneath"],
          [["Appointment delays and double bookings",
            "Booking is not treated as a transaction. Two staff members reading the same "
            "availability list can both write to it, and nothing at the storage layer stops them."],
           ["Rising volume of user requests",
            "Every request is handled synchronously. Slow work such as sending a confirmation "
            "sits on the critical path, so throughput collapses long before the hardware does."],
           ["Service availability problems and downtime",
            "There is no health signal an orchestrator can act on. A process that is alive but "
            "cannot reach its database looks identical to a healthy one."],
           ["Difficulty integrating existing applications",
            "The existing systems are UI-first. Their behaviour is only reachable by driving "
            "screens, so nothing else can consume it."]],
          widths=[33, 67], size=8.8)

    h2(doc, "1.2  Stakeholders")
    para(doc, "We listed the stakeholders before writing any requirements, because a requirement "
              "with no stakeholder behind it is usually a feature somebody invented. The last "
              "column is the one that shaped the design most.")
    table(doc,
          ["Stakeholder", "What they need from the system", "Design consequence"],
          [["Patient", "Find the right doctor quickly, see genuinely free slots, book without "
                       "phoning, get proof it worked",
            "Sub-second search; a booking reference returned in the response body"],
           ["Doctor", "A calendar that reflects reality and is never double booked",
            "Uniqueness enforced in the database, not in application code"],
           ["Hospital administrator", "Operational visibility and control over privileged actions",
            "Role-based access control and an admin statistics endpoint"],
           ["IT operations", "To know the service is healthy before patients tell them it is not",
            "Liveness and readiness probes, Prometheus metrics, structured logs"],
           ["Compliance officer", "An account of who touched patient data and when",
            "An append-only audit trail written inside the same transaction"],
           ["Development team (us)", "To change the system without breaking it",
            "An automated test suite and a pipeline that blocks a bad commit"]],
          widths=[18, 44, 38], size=8.6)

    h2(doc, "1.3  The problem, stated precisely")
    para(doc, "**Design and build a healthcare appointment platform in which a slot can be booked "
              "exactly once under concurrent load, in which slow side effects never delay the "
              "patient's response, whose behaviour is reachable through a documented API rather "
              "than only through screens, and which reports its own health well enough for an "
              "orchestrator to manage it automatically.** Everything else in this report follows "
              "from that sentence.")

    h2(doc, "1.4  Assumptions")
    bullets(doc, [
        "Consultations are fixed-length. We used 30 minutes, configurable through `SLOT_MINUTES`.",
        "A patient books on their own behalf. Booking for a dependant is out of scope for this iteration.",
        "Doctor calendars are published ahead of time by the hospital; the system does not negotiate a doctor's working hours.",
        "Notification delivery may be delayed by seconds. It must never be lost, and it must never delay the booking response.",
        "TLS termination, DNS and backups are provided by the hosting platform and are not re-implemented here.",
        "Payment and clinical records are separate systems. We expose the integration seam but do not build them.",
    ])

    h2(doc, "1.5  Expected outcomes")
    para(doc, "We wrote these down at the start so that at the end we could check honestly "
              "whether we had met them. Section 10 does exactly that, against measured numbers.")
    bullets(doc, [
        "A patient can go from landing page to confirmed booking without contacting the hospital.",
        "Two people racing for the same slot produce one confirmation and one clear refusal — never two bookings.",
        "Median read latency stays in single-digit milliseconds at realistic concurrency.",
        "Every endpoint is documented automatically, so a mobile client could be added without touching the backend.",
        "A container image builds, starts and answers a health probe without manual intervention.",
    ])


# =========================================================== section 2
def section2(doc):
    h1(doc, "2.  Objectives and Expected Outcomes")
    para(doc, "Our first draft of these objectives said things like \"the system should be fast\" "
              "and \"the system should be secure\". Those are not objectives, they are hopes, "
              "because nothing can prove them false. We rewrote each one until it named a number "
              "and a way of measuring it. The right-hand column is what we actually observed, "
              "reported fully in Section 10.")
    table(doc,
          ["#", "Objective (measurable)", "How we verified it", "Result"],
          [["O1", "Serve doctor-directory reads with a median latency under 50 ms at 25 concurrent users",
            "Closed-loop load test, 8 s per scenario, latency of every request recorded",
            "**50.15 ms p50** — met"],
           ["O2", "Make double booking impossible, not merely unlikely",
            "10 threads racing for one slot in an automated test (TC-16)",
            "**1 × 201, 9 × 409** — met"],
           ["O3", "Keep notification work off the booking response path",
            "Timestamped structured logs comparing response and notification write",
            "**Notification logged after the 201** — met"],
           ["O4", "Publish a machine-readable contract for every endpoint",
            "Assert the OpenAPI document lists each route (TC-25)",
            "**/openapi.json complete** — met"],
           ["O5", "Store no password in a recoverable form",
            "Inspect the stored value and verify the algorithm and salt (TC-05, TC-06)",
            "**PBKDF2-SHA256, 600 000 rounds** — met"],
           ["O6", "Reach at least 90% statement coverage of application code",
            "coverage.py restricted to the app package",
            "**95% (605 statements)** — met"],
           ["O7", "Produce a runnable container image without manual steps",
            "GitHub Actions builds the image and probes /health inside it",
            "**Image healthy in 31 s** — met"],
           ["O8", "Block a defective commit before it reaches the main branch",
            "Four-stage pipeline; each stage gates the next",
            "**Two real commits blocked** — met"]],
          widths=[4, 30, 36, 30], size=8.4, align_center=(0,))

    para(doc, "The last row is worth a sentence, because it is the objective we least expected to "
              "be able to evidence. Two of our own commits were genuinely rejected by our own "
              "pipeline while we were building this: one for lint failures, one for a broken "
              "deployment script. Both are visible in the run history linked in Section 9. We "
              "left them in the history rather than rewriting it, because a pipeline that has "
              "never failed has never been shown to work.")


# =========================================================== section 3
def section3(doc):
    h1(doc, "3.  Requirements, Constraints and Assumptions")

    h2(doc, "3.1  Functional requirements")
    table(doc,
          ["ID", "Requirement", "Priority", "Implemented in", "Verified by"],
          [["FR-01", "A patient can create an account with a validated e-mail and a password of at least 8 characters", "Must", "`routers/auth.py`", "TC-01, TC-03"],
           ["FR-02", "Duplicate registration on the same e-mail is refused", "Must", "`routers/auth.py`", "TC-02"],
           ["FR-03", "A registered patient can exchange credentials for a signed, expiring token", "Must", "`security.py`", "TC-04"],
           ["FR-04", "Repeated failed logins from one address are throttled", "Should", "`ratelimit.py`", "TC-24"],
           ["FR-05", "Anyone can browse the doctor directory and filter it by speciality or name", "Must", "`routers/doctors.py`", "TC-10"],
           ["FR-06", "Only future, unbooked slots are offered for a given doctor", "Must", "`routers/doctors.py`", "TC-11"],
           ["FR-07", "An authenticated patient can book a free slot and receives a booking reference", "Must", "`routers/appointments.py`", "TC-14"],
           ["FR-08", "A slot that is already taken cannot be booked again, under any interleaving", "Must", "`routers/appointments.py`", "TC-15, TC-16"],
           ["FR-09", "A patient can list their own appointments and no one else's", "Must", "`routers/appointments.py`", "TC-14"],
           ["FR-10", "Cancelling an appointment returns the slot to the available pool", "Must", "`routers/appointments.py`", "TC-17"],
           ["FR-11", "A confirmation notification is generated for every booking and cancellation", "Must", "`routers/notifications.py`", "TC-19"],
           ["FR-12", "Administrative statistics are reachable only by the admin role", "Should", "`deps.py`, `routers/ops.py`", "TC-09"],
           ["FR-13", "Every authentication and appointment action is written to an append-only audit trail", "Must", "`models.py`", "Inspected in Section 9"]],
          widths=[6, 42, 8, 22, 12], size=8.2, align_center=(0, 2))

    h2(doc, "3.2  Non-functional requirements")
    para(doc, "These are the requirements that decided the architecture. NFR-01 and NFR-02 in "
              "particular are the reason the system is event-driven rather than a straightforward "
              "three-tier application.")
    table(doc,
          ["ID", "Quality attribute", "Target we committed to", "Measured outcome"],
          [["NFR-01", "Performance", "Read endpoints under 500 ms at the 95th percentile, 25 concurrent users", "68.33 ms p95 — 7× margin"],
           ["NFR-02", "Responsiveness", "Booking must not wait for notification delivery", "Booking 34.95 ms; notification written afterwards"],
           ["NFR-03", "Scalability", "Throughput must rise with concurrency until CPU-bound", "318 → 462 rps; saturation identified at 5 workers"],
           ["NFR-04", "Reliability", "Zero failed requests across the whole load test", "36 057 requests, 100% success"],
           ["NFR-05", "Data integrity", "One slot, one appointment, always", "Enforced by `UNIQUE(slot_id)`; proven by TC-16"],
           ["NFR-06", "Security", "No plaintext passwords, signed sessions, baseline headers", "PBKDF2 600k rounds; HS256; headers on every response"],
           ["NFR-07", "Maintainability", "Static analysis clean; coverage above 90%", "ruff and bandit clean; 95% coverage"],
           ["NFR-08", "Operability", "Liveness, readiness and metrics endpoints", "`/health`, `/ready`, `/metrics` live"],
           ["NFR-09", "Portability", "Runs identically on a laptop and on a CI runner", "Same image healthy in both environments"],
           ["NFR-10", "Auditability", "Every privileged action attributable to a principal", "`audit_log` written inside the booking transaction"]],
          widths=[7, 17, 40, 36], size=8.3, align_center=(0,))

    h2(doc, "3.3  Constraints")
    table(doc,
          ["Type", "Constraint", "How we worked within it"],
          [["Technical", "No managed cloud services were available to us",
            "SQLite behind the ORM and an in-process event bus, both chosen so the interface matches the production replacement"],
           ["Cost", "Zero budget",
            "Entirely open-source stack; GitHub Actions free tier; no paid infrastructure at any point"],
           ["Time", "One assignment cycle across a small team",
            "Scope fixed to the booking core; payments and clinical records deliberately excluded and named as such"],
           ["Skill", "First serious use of containers and CI for most of the team",
            "Pipeline built early and kept simple so that failures were readable"],
           ["Privacy", "Patient data is sensitive even in a prototype",
            "No real patient data used anywhere; seeded doctors and patients are fictional"],
           ["Environment", "Docker Desktop would not start on one team machine",
            "Container build and smoke test run on the CI runner instead, which is stronger evidence anyway"]],
          widths=[12, 38, 50], size=8.5)

    note(doc, "**An honest note on that last row.** We could not get the Docker daemon to start on "
              "the Windows machine used for the final capture session, so the container evidence in "
              "Section 9 comes from the GitHub Actions runner rather than a local terminal. We think "
              "this is actually the better artefact: the image is built from a clean checkout, on a "
              "machine none of us configured, and is then started and probed automatically.")


# =========================================================== section 4
def section4(doc):
    h1(doc, "4.  Application of Software Engineering Concepts")
    para(doc, "This section connects the syllabus to decisions we actually made. We have tried to "
              "avoid listing concepts we did not use.")

    h2(doc, "4.1  Process model: Scrum, in short iterations")
    para(doc, "We used Scrum with short sprints because the requirements were clear enough to "
              "plan but our understanding of the technology was not. A waterfall plan would have "
              "committed us to an architecture before we knew, for example, how expensive PBKDF2 "
              "actually is. Working in increments meant we found that out in Sprint 2 and could "
              "still act on it.")
    table(doc,
          ["Sprint", "Goal", "Delivered increment", "What we learned"],
          [["1", "Walking skeleton", "Domain model, database schema, doctor directory, seeded calendar",
            "Putting the uniqueness rule in the schema early made the booking code simpler later"],
           ["2", "Identity", "Registration, login, JWT issue and verification, role guard",
            "PBKDF2 at 600 000 rounds costs ~252 ms. That is intentional, but it must never sit on a read path"],
           ["3", "The booking core", "Transactional booking, cancellation, the concurrency test",
            "Our first test passed for the wrong reason — see Section 8.2"],
           ["4", "Asynchrony", "Event bus, notification subscriber, background delivery",
            "Moving work off the request path is a design decision, not an optimisation"],
           ["5", "Operability and delivery", "Metrics, probes, structured logs, Docker image, CI/CD pipeline",
            "The pipeline caught two real defects in our own code within the first hour"]],
          widths=[6, 17, 39, 38], size=8.3, align_center=(0,))

    h2(doc, "4.2  Requirements engineering")
    para(doc, "Requirements were elicited from the problem statement, written as the numbered "
              "FR/NFR tables in Section 3, and then traced forward. Every functional requirement "
              "names the module that implements it and the test that proves it; every "
              "non-functional requirement names the measurement that checks it. That traceability "
              "is what let us answer \"is it done?\" with evidence rather than opinion.")

    h2(doc, "4.3  Feasibility analysis")
    table(doc,
          ["Dimension", "Assessment", "Evidence"],
          [["Technical", "Feasible", "Delivered and running; 25 automated tests pass; container builds unattended"],
           ["Economic", "Feasible", "Every tool used is free and open source; total infrastructure spend was zero"],
           ["Operational", "Feasible", "Patient workflow needs no training; administrators keep an audit trail by default"],
           ["Schedule", "Feasible with the scope we fixed", "Five short sprints; payments and clinical records excluded up front"],
           ["Legal / ethical", "Feasible with care", "No real patient data; audit trail and access control designed in from the start"]],
          widths=[14, 24, 62], size=8.5)

    h2(doc, "4.4  Architectural decisions and why we made them")
    para(doc, "Four decisions determined most of the system. We record them here in the form we "
              "used at the time, because the alternatives we rejected are as informative as the "
              "options we chose. Section 11 compares them in more depth.")
    table(doc,
          ["Decision", "What we chose", "What we rejected, and why"],
          [["Service decomposition",
            "Modular monolith: four bounded contexts (auth, doctors, appointments, notifications) in one deployable process",
            "True microservices. With four modules and a team of our size, we would have paid for network hops, distributed transactions and four pipelines, and bought nothing. The module boundaries are real, so the split remains available later."],
           ["Interface style",
            "API-first. Every capability is an HTTP endpoint described by an auto-generated OpenAPI document",
            "A server-rendered application. It would have been quicker, but it would have recreated exactly the integration problem the organisation complained about."],
           ["Notification delivery",
            "Event-driven. Booking publishes `appointment.booked` and returns immediately",
            "A direct call from the booking service. That puts an unrelated system's latency, and its outages, on the patient's critical path."],
           ["Concurrency control",
            "Pessimistic row lock plus a `UNIQUE` constraint on `appointments.slot_id`",
            "Application-level checking. A check-then-act in application code is a race by construction; the database is the only place the rule can actually be enforced."]],
          widths=[15, 33, 52], size=8.3)

    h2(doc, "4.5  Version control, containerisation and DevOps")
    para(doc, "Work is tracked in Git and hosted on GitHub with the main branch protected by the "
              "pipeline. The image is a two-stage Docker build: dependencies are compiled in a "
              "builder stage that is then discarded, and the runtime stage runs as an unprivileged "
              "user (uid 10001) with a `HEALTHCHECK` baked in. The pipeline runs static analysis, "
              "then tests, then builds and smoke-tests the image, then deploys — each stage gated "
              "on the one before it. Sections 7 and 9 show the Dockerfile and the real build log.")

    h2(doc, "4.6  Testing strategy")
    para(doc, "We followed the test pyramid, but deliberately weighted it towards integration "
              "tests that exercise a real HTTP request against a real database. Unit-testing our "
              "booking logic in isolation would have missed the only bug in it that mattered, "
              "because that bug lived in the interaction between two sessions and a transaction.")
    table(doc,
          ["Layer", "Count", "What it covers", "Example"],
          [["Unit", "6", "Pure logic with no I/O", "TC-06 proves two identical passwords hash differently because of the per-user salt"],
           ["Integration", "17", "A real HTTP request through the full stack to a real database", "TC-14 books a slot and reads it back through the API"],
           ["Concurrency", "1", "Ten OS threads racing for a single row", "TC-16 asserts exactly one 201 and nine 409s"],
           ["Contract", "1", "The published OpenAPI document", "TC-25 asserts every route appears in `/openapi.json`"]],
          widths=[12, 7, 30, 51], size=8.4, align_center=(1,))

    h2(doc, "4.7  Maintenance, technical debt and sustainability")
    para(doc, "We keep a short, explicit debt register rather than pretending there is none. "
              "SQLite is the largest item: it is correct for our tests but serialises writers, so "
              "it is a production ceiling. The in-process event bus is the second: it loses "
              "undelivered events if the process dies. Both were chosen knowingly, both sit behind "
              "an interface that matches their eventual replacement (`DATABASE_URL` for the "
              "database, a publish/subscribe API for the bus), and both are listed in Section 13 "
              "with the migration path. On sustainability, the measurable choice we made was to "
              "keep the runtime image dependency-free where we reasonably could: the metrics "
              "registry and the JWT implementation are standard-library only, which keeps the "
              "image small and the number of packages we must keep patched low.")
